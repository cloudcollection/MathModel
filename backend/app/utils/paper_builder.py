from __future__ import annotations

import re
from pathlib import Path


def write_paper_files(markdown: str, output_dir: Path) -> tuple[Path, Path | None]:
    """Write result.md and, when python-docx is installed, result.docx."""

    output_dir.mkdir(parents=True, exist_ok=True)
    md_path = output_dir / "result.md"
    md_path.write_text(markdown, encoding="utf-8")

    docx_path = output_dir / "result.docx"
    try:
        from docx import Document
    except Exception:
        return md_path, None

    document = Document()
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(_clean_markdown(line[2:]), level=0)
        elif line.startswith("## "):
            document.add_heading(_clean_markdown(line[3:]), level=1)
        elif line.startswith("### "):
            document.add_heading(_clean_markdown(line[4:]), level=2)
        elif line.startswith("- "):
            document.add_paragraph(_clean_markdown(line[2:]), style="List Bullet")
        else:
            document.add_paragraph(_clean_markdown(line))

    document.save(docx_path)
    return md_path, docx_path


def _clean_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text

